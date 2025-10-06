"""
LlamaIndex Tools Implementation for Forge1

Production-ready tools that integrate LlamaIndex capabilities with Forge1's
security, tenancy, and observability systems. Each tool enforces RBAC,
applies DLP redaction, and tracks usage metrics.
"""

import asyncio
import logging
import os
import tempfile
import uuid
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import json
import base64

# Document processing imports
import pypdf
from docx import Document as DocxDocument
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None

# LlamaIndex imports
from llama_index.core.schema import Document, TextNode
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.readers.file import PDFReader, DocxReader

# Google Drive imports
try:
    from googleapiclient.discovery import build
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    GOOGLE_DRIVE_AVAILABLE = True
except ImportError:
    GOOGLE_DRIVE_AVAILABLE = False

# Slack imports
try:
    from slack_sdk.web.async_client import AsyncWebClient
    from slack_sdk.errors import SlackApiError
    SLACK_AVAILABLE = True
except ImportError:
    SLACK_AVAILABLE = False

# Forge1 imports
from forge1.integrations.llamaindex_adapter import (
    LlamaIndexTool, ToolType, ExecutionContext, ToolResult
)
from forge1.core.memory_manager import MemoryManager, MemoryQuery, MemoryType
from forge1.core.model_router import ModelRouter
from forge1.integrations.llamaindex_model_shim import ModelShimFactory

logger = logging.getLogger(__name__)

class DocumentParserTool(LlamaIndexTool):
    """Document parser tool with OCR fallback capability"""
    
    def __init__(self, model_router: ModelRouter, memory_manager: MemoryManager, **kwargs):
        super().__init__(tool_type=ToolType.DOCUMENT_PARSER, **kwargs)
        self.model_router = model_router
        self.memory_manager = memory_manager
        self.node_parser = SimpleNodeParser()
        
        # Initialize model shim factory
        self.model_shim_factory = ModelShimFactory(
            model_router=model_router,
            audit_logger=self.audit_logger
        )
    
    async def acall(self, context: ExecutionContext, **kwargs) -> Dict[str, Any]:
        """Parse document with OCR fallback"""
        try:
            # Enforce RBAC
            await self._enforce_rbac(context, "document:read")
            
            # Extract parameters
            document_path = kwargs.get("document_path")
            document_content = kwargs.get("document_content")  # Base64 encoded
            document_format = kwargs.get("document_format", "auto")
            use_ocr = kwargs.get("use_ocr", True)
            
            if not document_path and not document_content:
                raise ValueError("Either document_path or document_content must be provided")
            
            # Handle document content
            if document_content:
                # Decode base64 content and save to temp file
                content_bytes = base64.b64decode(document_content)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{document_format}")
                temp_file.write(content_bytes)
                temp_file.close()
                document_path = temp_file.name
            
            # Determine document format
            if document_format == "auto":
                document_format = self._detect_format(document_path)
            
            # Parse document
            parsed_content = await self._parse_document(
                document_path, document_format, use_ocr, context
            )
            
            # Clean up temp file if created
            if document_content and os.path.exists(document_path):
                os.unlink(document_path)
            
            # Store parsed content in memory for future reference
            await self._store_parsed_content(parsed_content, context)
            
            return {
                "success": True,
                "parsed_content": parsed_content,
                "document_format": document_format,
                "node_count": len(parsed_content.get("nodes", [])),
                "text_length": len(parsed_content.get("text", ""))
            }
            
        except Exception as e:
            logger.error(f"Document parsing failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _detect_format(self, file_path: str) -> str:
        """Detect document format from file extension"""
        suffix = Path(file_path).suffix.lower()
        format_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".txt": "txt",
            ".png": "image",
            ".jpg": "image",
            ".jpeg": "image"
        }
        return format_map.get(suffix, "unknown")
    
    async def _parse_document(
        self, 
        file_path: str, 
        doc_format: str, 
        use_ocr: bool,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """Parse document based on format"""
        
        if doc_format == "pdf":
            return await self._parse_pdf(file_path, use_ocr, context)
        elif doc_format == "docx":
            return await self._parse_docx(file_path, context)
        elif doc_format == "txt":
            return await self._parse_text(file_path, context)
        elif doc_format == "image" and use_ocr:
            return await self._parse_image_ocr(file_path, context)
        else:
            raise ValueError(f"Unsupported document format: {doc_format}")
    
    async def _parse_pdf(self, file_path: str, use_ocr: bool, context: ExecutionContext) -> Dict[str, Any]:
        """Parse PDF document with OCR fallback"""
        try:
            # Try standard PDF parsing first
            reader = PDFReader()
            documents = reader.load_data(file_path)
            
            if documents and documents[0].text.strip():
                # Successfully extracted text
                text = "\n".join([doc.text for doc in documents])
                nodes = self.node_parser.get_nodes_from_documents(documents)
                
                return {
                    "text": text,
                    "nodes": [{"text": node.text, "metadata": node.metadata} for node in nodes],
                    "extraction_method": "standard_pdf",
                    "page_count": len(documents)
                }
            
            elif use_ocr and OCR_AVAILABLE:
                # Fallback to OCR
                logger.info(f"PDF text extraction failed, falling back to OCR for {file_path}")
                return await self._parse_pdf_ocr(file_path, context)
            
            else:
                raise ValueError("PDF text extraction failed and OCR not available")
                
        except Exception as e:
            if use_ocr and OCR_AVAILABLE:
                logger.warning(f"Standard PDF parsing failed, trying OCR: {e}")
                return await self._parse_pdf_ocr(file_path, context)
            else:
                raise e
    
    async def _parse_pdf_ocr(self, file_path: str, context: ExecutionContext) -> Dict[str, Any]:
        """Parse PDF using OCR"""
        if not OCR_AVAILABLE:
            raise ValueError("OCR not available - install pytesseract and PIL")
        
        try:
            # Convert PDF pages to images and OCR each page
            import fitz  # PyMuPDF for PDF to image conversion
            
            doc = fitz.open(file_path)
            extracted_text = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("ppm")
                
                # Convert to PIL Image
                from io import BytesIO
                img = Image.open(BytesIO(img_data))
                
                # OCR the image
                page_text = pytesseract.image_to_string(img)
                extracted_text.append(page_text)
            
            doc.close()
            
            # Combine all text
            full_text = "\n".join(extracted_text)
            
            # Create document and nodes
            document = Document(text=full_text)
            nodes = self.node_parser.get_nodes_from_documents([document])
            
            return {
                "text": full_text,
                "nodes": [{"text": node.text, "metadata": node.metadata} for node in nodes],
                "extraction_method": "ocr",
                "page_count": len(extracted_text)
            }
            
        except ImportError:
            raise ValueError("PyMuPDF (fitz) required for PDF OCR - install with: pip install PyMuPDF")
        except Exception as e:
            raise ValueError(f"OCR extraction failed: {e}")
    
    async def _parse_docx(self, file_path: str, context: ExecutionContext) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            reader = DocxReader()
            documents = reader.load_data(file_path)
            
            text = "\n".join([doc.text for doc in documents])
            nodes = self.node_parser.get_nodes_from_documents(documents)
            
            return {
                "text": text,
                "nodes": [{"text": node.text, "metadata": node.metadata} for node in nodes],
                "extraction_method": "docx_reader",
                "document_count": len(documents)
            }
            
        except Exception as e:
            raise ValueError(f"DOCX parsing failed: {e}")
    
    async def _parse_text(self, file_path: str, context: ExecutionContext) -> Dict[str, Any]:
        """Parse plain text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            document = Document(text=text)
            nodes = self.node_parser.get_nodes_from_documents([document])
            
            return {
                "text": text,
                "nodes": [{"text": node.text, "metadata": node.metadata} for node in nodes],
                "extraction_method": "text_reader",
                "character_count": len(text)
            }
            
        except Exception as e:
            raise ValueError(f"Text parsing failed: {e}")
    
    async def _parse_image_ocr(self, file_path: str, context: ExecutionContext) -> Dict[str, Any]:
        """Parse image using OCR"""
        if not OCR_AVAILABLE:
            raise ValueError("OCR not available - install pytesseract and PIL")
        
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            
            document = Document(text=text)
            nodes = self.node_parser.get_nodes_from_documents([document])
            
            return {
                "text": text,
                "nodes": [{"text": node.text, "metadata": node.metadata} for node in nodes],
                "extraction_method": "image_ocr",
                "image_size": img.size
            }
            
        except Exception as e:
            raise ValueError(f"Image OCR failed: {e}")
    
    async def _store_parsed_content(self, parsed_content: Dict[str, Any], context: ExecutionContext) -> None:
        """Store parsed content in Forge1 memory system"""
        try:
            # Store the parsed document content
            await self.memory_manager.store_memory(
                content=parsed_content,
                memory_type=MemoryType.DOCUMENT,
                metadata={
                    "tool_type": "document_parser",
                    "tenant_id": context.tenant_id,
                    "employee_id": context.employee_id,
                    "case_id": context.case_id,
                    "extraction_method": parsed_content.get("extraction_method"),
                    "node_count": len(parsed_content.get("nodes", []))
                }
            )
            
        except Exception as e:
            logger.warning(f"Failed to store parsed content in memory: {e}")

class KBSearchTool(LlamaIndexTool):
    """Knowledge base search tool with vector/hybrid retrieval"""
    
    def __init__(self, model_router: ModelRouter, memory_manager: MemoryManager, **kwargs):
        super().__init__(tool_type=ToolType.KB_SEARCH, **kwargs)
        self.model_router = model_router
        self.memory_manager = memory_manager
        
        # Initialize model shim factory
        self.model_shim_factory = ModelShimFactory(
            model_router=model_router,
            audit_logger=self.audit_logger
        )
    
    async def acall(self, context: ExecutionContext, **kwargs) -> Dict[str, Any]:
        """Search knowledge base using vector/hybrid retrieval"""
        try:
            # Enforce RBAC
            await self._enforce_rbac(context, "knowledge_base:search")
            
            # Extract parameters
            query = kwargs.get("query", "")
            max_results = kwargs.get("max_results", 10)
            similarity_threshold = kwargs.get("similarity_threshold", 0.7)
            search_types = kwargs.get("search_types", [MemoryType.DOCUMENT, MemoryType.CONVERSATION])
            
            if not query:
                raise ValueError("Query parameter is required")
            
            # Perform search using Forge1's memory manager
            search_query = MemoryQuery(
                query_text=query,
                memory_types=search_types,
                limit=max_results,
                min_relevance_score=similarity_threshold,
                employee_ids=[context.employee_id],  # Scope to current employee
                max_age_days=365  # Search within last year
            )
            
            search_response = await self.memory_manager.search_memories(
                query=search_query,
                user_id=context.employee_id
            )
            
            # Format results
            results = []
            for result in search_response.results:
                results.append({
                    "memory_id": result.memory.id,
                    "content": result.memory.content,
                    "summary": result.memory.summary,
                    "similarity_score": result.similarity_score,
                    "memory_type": result.memory.memory_type.value,
                    "created_at": result.memory.created_at.isoformat(),
                    "metadata": result.memory.metadata if hasattr(result.memory, 'metadata') else {}
                })
            
            return {
                "success": True,
                "query": query,
                "results": results,
                "total_count": search_response.total_count,
                "query_time_ms": search_response.query_time_ms
            }
            
        except Exception as e:
            logger.error(f"Knowledge base search failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "query": kwargs.get("query", "")
            }

class DriveFetchTool(LlamaIndexTool):
    """Google Drive fetch tool with read-only access"""
    
    def __init__(self, model_router: ModelRouter, memory_manager: MemoryManager, **kwargs):
        super().__init__(tool_type=ToolType.DRIVE_FETCH, **kwargs)
        self.model_router = model_router
        self.memory_manager = memory_manager
    
    async def acall(self, context: ExecutionContext, **kwargs) -> Dict[str, Any]:
        """Fetch document from Google Drive"""
        try:
            # Enforce RBAC
            await self._enforce_rbac(context, "drive:read")
            
            if not GOOGLE_DRIVE_AVAILABLE:
                raise ValueError("Google Drive integration not available - install google-api-python-client")
            
            # Extract parameters
            file_id = kwargs.get("file_id")
            file_path = kwargs.get("file_path")  # Alternative: path-based lookup
            
            if not file_id and not file_path:
                raise ValueError("Either file_id or file_path must be provided")
            
            # Get tenant-scoped Google Drive credentials
            credentials_json = await self._get_tenant_secret("google_drive_credentials", context)
            
            # Initialize Drive service
            service = await self._create_drive_service(credentials_json)
            
            # Fetch file
            if file_id:
                file_data = await self._fetch_file_by_id(service, file_id)
            else:
                file_data = await self._fetch_file_by_path(service, file_path)
            
            # Store access log
            await self._log_drive_access(context, file_data)
            
            return {
                "success": True,
                "file_data": file_data,
                "access_time": datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            logger.error(f"Google Drive fetch failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _create_drive_service(self, credentials_json: str):
        """Create Google Drive service with tenant credentials"""
        try:
            import json
            creds_dict = json.loads(credentials_json)
            
            # Create credentials object
            creds = Credentials.from_authorized_user_info(creds_dict)
            
            # Build service
            service = build('drive', 'v3', credentials=creds)
            return service
            
        except Exception as e:
            raise ValueError(f"Failed to create Drive service: {e}")
    
    async def _fetch_file_by_id(self, service, file_id: str) -> Dict[str, Any]:
        """Fetch file by Google Drive file ID"""
        try:
            # Get file metadata
            file_metadata = service.files().get(fileId=file_id).execute()
            
            # Get file content
            file_content = service.files().get_media(fileId=file_id).execute()
            
            return {
                "file_id": file_id,
                "name": file_metadata.get("name"),
                "mime_type": file_metadata.get("mimeType"),
                "size": file_metadata.get("size"),
                "modified_time": file_metadata.get("modifiedTime"),
                "content": base64.b64encode(file_content).decode('utf-8'),
                "metadata": file_metadata
            }
            
        except Exception as e:
            raise ValueError(f"Failed to fetch file {file_id}: {e}")
    
    async def _fetch_file_by_path(self, service, file_path: str) -> Dict[str, Any]:
        """Fetch file by path (search by name)"""
        try:
            # Search for file by name
            query = f"name='{os.path.basename(file_path)}'"
            results = service.files().list(q=query).execute()
            
            files = results.get('files', [])
            if not files:
                raise ValueError(f"File not found: {file_path}")
            
            # Use first matching file
            file_id = files[0]['id']
            return await self._fetch_file_by_id(service, file_id)
            
        except Exception as e:
            raise ValueError(f"Failed to fetch file by path {file_path}: {e}")
    
    async def _log_drive_access(self, context: ExecutionContext, file_data: Dict[str, Any]) -> None:
        """Log Google Drive access for audit"""
        try:
            await self.audit_logger.log_access_event(
                event_type="drive_access",
                user_id=context.employee_id,
                tenant_id=context.tenant_id,
                details={
                    "file_id": file_data.get("file_id"),
                    "file_name": file_data.get("name"),
                    "file_size": file_data.get("size"),
                    "case_id": context.case_id,
                    "request_id": context.request_id
                }
            )
        except Exception as e:
            logger.warning(f"Failed to log Drive access: {e}")

class SlackPostTool(LlamaIndexTool):
    """Slack post tool for tenant-scoped channels"""
    
    def __init__(self, model_router: ModelRouter, memory_manager: MemoryManager, **kwargs):
        super().__init__(tool_type=ToolType.SLACK_POST, **kwargs)
        self.model_router = model_router
        self.memory_manager = memory_manager
    
    async def acall(self, context: ExecutionContext, **kwargs) -> Dict[str, Any]:
        """Post message to Slack channel"""
        try:
            # Enforce RBAC
            await self._enforce_rbac(context, "slack:post")
            
            if not SLACK_AVAILABLE:
                raise ValueError("Slack integration not available - install slack-sdk")
            
            # Extract parameters
            channel = kwargs.get("channel")
            message = kwargs.get("message", "")
            attachments = kwargs.get("attachments", [])
            thread_ts = kwargs.get("thread_ts")  # For threaded replies
            
            if not channel or not message:
                raise ValueError("Both channel and message parameters are required")
            
            # Get tenant-scoped Slack token
            slack_token = await self._get_tenant_secret("slack_bot_token", context)
            
            # Initialize Slack client
            client = AsyncWebClient(token=slack_token)
            
            # Post message
            response = await client.chat_postMessage(
                channel=channel,
                text=message,
                attachments=attachments,
                thread_ts=thread_ts
            )
            
            # Log the post
            await self._log_slack_post(context, channel, message, response)
            
            return {
                "success": True,
                "message_ts": response["ts"],
                "channel": response["channel"],
                "message": message,
                "post_time": datetime.now(timezone.utc).isoformat()
            }
            
        except SlackApiError as e:
            logger.error(f"Slack API error: {e}")
            return {
                "success": False,
                "error": f"Slack API error: {e.response['error']}"
            }
        except Exception as e:
            logger.error(f"Slack post failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _log_slack_post(
        self, 
        context: ExecutionContext, 
        channel: str, 
        message: str, 
        response: Dict[str, Any]
    ) -> None:
        """Log Slack post for audit"""
        try:
            await self.audit_logger.log_access_event(
                event_type="slack_post",
                user_id=context.employee_id,
                tenant_id=context.tenant_id,
                details={
                    "channel": channel,
                    "message_length": len(message),
                    "message_ts": response.get("ts"),
                    "case_id": context.case_id,
                    "request_id": context.request_id
                }
            )
        except Exception as e:
            logger.warning(f"Failed to log Slack post: {e}")